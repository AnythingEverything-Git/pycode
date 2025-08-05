import os
import json
import re

import graphviz
from groq import Groq

from docx import Document
from docx.shared import Inches
from markdown import markdown
from bs4 import BeautifulSoup
from docx2pdf import convert

# OpenAI API Key (replace with env variable in production)
# client = OpenAI(api_key="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")


# ---------------------
# CONFIG
# ---------------------
# BRD_FILE = "business_requirements.txt"  # BRD in text format
# MODEL_NAME = "mistral"  # Ollama model (mistral is faster for JSON)
# CHUNK_SIZE = 500  # Words per chunk
# OLLAMA_PATH = r"C:\Users\SOUMI\AppData\Local\Programs\Ollama\ollama.exe"

BRD_FILE = "business_requirements.txt"
CHUNK_SIZE = 2000  # words per chunk
MODEL_NAME = "llama3-8b-8192"  # Groq LLaMA model

client = Groq(api_key="")


# ---------------------
# STEP 1: Load and Chunk BRD
# ---------------------
def load_brd(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()

def split_document(text, max_words=CHUNK_SIZE):
    words = text.split()
    return [" ".join(words[i:i + max_words]) for i in range(0, len(words), max_words)]


def fix_ai_json(raw_output: str):
    """
    Cleans and repairs AI JSON output to handle:
    1. Missing {} or []
    2. Missing last closing bracket/brace
    3. Single → double quote conversion (safe)
    4. Trailing commas
    """
    if not raw_output:
        return None

    # 1️⃣ Extract first JSON block: { ... } or [ ... ]
    match = re.search(r'(\{.*|\[.*)', raw_output, re.DOTALL)
    json_str = match.group(0).strip() if match else raw_output.strip()

    # 2️⃣ Remove trailing commas before } or ]
    json_str = re.sub(r',(\s*[}\]])', r'\1', json_str)

    # 3️⃣ Safely convert Python-like dict to JSON
    # Convert only keys/values in single quotes to double quotes without touching already double-quoted strings
    def safe_quote_replace(match):
        return '"' + match.group(1) + '"'

    # Replace single-quoted keys
    json_str = re.sub(r"(?<!\")'([A-Za-z0-9_ ]+)'(?=\s*:)", safe_quote_replace, json_str)
    # Replace single-quoted string values
    json_str = re.sub(r":\s*'([^']*)'", lambda m: ':"{}"'.format(m.group(1)), json_str)

    # 4️⃣ Auto-complete missing closing braces/brackets
    open_curly = json_str.count("{")
    close_curly = json_str.count("}")
    if close_curly < open_curly:
        json_str += "}" * (open_curly - close_curly)

    open_square = json_str.count("[")
    close_square = json_str.count("]")
    if close_square < open_square:
        json_str += "]" * (open_square - close_square)

    # 5️⃣ Attempt JSON parsing
    try:
        return json.loads(json_str)
    except json.JSONDecodeError:
        return None


def generate_ai_layout(input_json: dict):
    prompt = f"""
    You are a senior software architect and C4 diagram expert.

    I will provide a merged system architecture JSON containing:
    - "actors": external users or systems
    - "microservices": internal services (with optional DBs)
    - "databases": database components
    - "events": interactions between actors, services, and databases

    Here is the architecture JSON:
    {json.dumps(input_json, indent=2)}

    Your task:
    Generate a **production-ready HLD JSON** for a C4-style Container Diagram, reflecting a professional microservices architecture with inter-service interactions like this reference pattern:

        End User -> API Gateway -> Microservices Layer -> External Systems
        Microservices communicate with each other using REST, gRPC, Events (Kafka/RabbitMQ)
        Databases align below their respective services
        Analytics/Logging services are optional downstream services

    The output JSON must have this structure:

    {{
  "nodes": [
    {{
      "name": "string - node name (actor/service/db/external)",
      "type": "actor | gateway | service | db | external",
      "layer": "frontend | gateway | microservice | data | external",
      "x": int,
      "y": int,
      "color": "lightgreen | lightblue | lightyellow | lightgray"
    }}
  ],
  "edges": [
    {{
      "from": "string - source node name",
      "to": "string - target node name",
      "label": "REST API | gRPC Call | Message Queue | DB Query | Event Stream",
      "direction": "uni | bi",
      "criticality": "High | Medium | Low"
    }}
  ],
  "microservice_interactions": [
    {{
      "caller": "string - microservice name",
      "callee": "string - microservice name",
      "protocol": "REST | gRPC | Event",
      "sync": true,
      "purpose": "string - short description"
    }}
  ]
}}

    Layout & Styling Rules:
    1. Place **End Users** (actors) at the top (y=0)
    2. Place **API Gateway or BFF services** in the second row (y=-1)
    3. Place **microservices layer** in the middle (y=-2)
    4. Place **databases** below their microservices (y=-3)
    5. Place **external systems** at the bottom (y=-4)
    6. Maintain horizontal spacing of 2 units between nodes in each layer
    7. Use these colors:
       - Actors → lightgreen
       - Gateways → lightskyblue
       - Internal services → lightblue
       - Databases → lightyellow
       - External systems → lightgray
    8. Interactions:
       - REST → blue solid arrow
       - gRPC → cyan solid arrow
       - Async Event/Queue → green dashed arrow
       - DB Query → brown dashed arrow
    9. Explicitly include **microservice-to-microservice interactions** in `edges` and `microservice_interactions`
    10. Ensure JSON is **valid**, no nulls, no trailing commas, properly closed braces/brackets
    11. Output **only valid JSON**, no commentary or markdown
    """

    response = ai_response(prompt, "You are a helpful assistant that outputs JSON only.")

    raw_output = response.choices[0].message.content.strip()
    return raw_output



def render_layout_to_png(layout_json: dict, output_file="c4_ai_layout"):
    """
    Renders the AI-generated layout JSON to a professional HLD diagram (PNG) using Graphviz.
    Surrounds all microservices with a dotted-line cluster.
    """
    dot = graphviz.Digraph(format='png')
    dot.attr(rankdir='TB', splines='ortho', fontname='Helvetica')

    # ----------------------
    # Step 1: Identify Microservices
    # ----------------------
    microservices = [node for node in layout_json.get("nodes", []) if node.get("type", "").lower() == "service"]
    other_nodes = [node for node in layout_json.get("nodes", []) if node.get("type", "").lower() != "service"]

    # ----------------------
    # Step 2: Render Non-Microservice Nodes
    # ----------------------
    for node in other_nodes:
        node_name = node.get("name", "").strip()
        node_type = node.get("type", "").lower()
        if not node_name or node_name.lower() == "null":
            continue

        # Professional color & shape mapping
        shape = "ellipse" if node_type == "actor" else "box"
        style = "filled"
        fillcolor = node.get("color", "white")

        if node_type == "db":
            shape = "cylinder"
        elif node_type == "external":
            style = "dashed,filled"

        dot.node(
            node_name,
            shape=shape,
            style=style,
            fillcolor=fillcolor,
            fontsize="12",
            fontcolor="black",
            width="1.2",
            height="0.8",
            pos=f'{node.get("x", 0)},{node.get("y", 0)}!'
        )

    # ----------------------
    # Step 3: Render Microservices in a Cluster
    # ----------------------
    if microservices:
        with dot.subgraph(name="cluster_microservices") as c:
            c.attr(label="Microservices Layer", color="gray", style="dotted", fontsize="14", fontname="Helvetica-Bold")
            for node in microservices:
                node_name = node.get("name", "").strip()
                if not node_name:
                    continue
                c.node(
                    node_name,
                    shape="box3d",
                    style="filled",
                    fillcolor=node.get("color", "lightblue"),
                    fontsize="12",
                    fontcolor="black",
                    width="1.2",
                    height="0.8",
                    pos=f'{node.get("x", 0)},{node.get("y", 0)}!'
                )

    # ----------------------
    # Step 4: Render Edges
    # ----------------------
    for edge in layout_json.get("edges", []):
        src = edge.get("from")
        dst = edge.get("to")
        label = edge.get("label", "").strip()

        if not src or not dst:
            continue

        # Determine color/style based on label
        edge_color = "black"
        edge_style = "solid"
        if "db" in label.lower():
            edge_color = "brown"; edge_style = "dashed"
        elif "rest" in label.lower() or "api" in label.lower():
            edge_color = "black"
        elif "queue" in label.lower() or "event" in label.lower():
            edge_color = "black"; edge_style = "dashed"
        elif "grpc" in label.lower():
            edge_color = "cyan"

        dot.edge(
            src,
            dst,
            xlabel=label,
            color=edge_color,
            fontcolor=edge_color,
            fontsize="10",
            style=edge_style,
            arrowsize="0.7"
        )

    # ----------------------
    # Step 5: Render PNG
    # ----------------------
    output_path = dot.render(output_file, cleanup=True)
    print(f"✅ Professional HLD PNG generated with microservices cluster: {output_path}")
    return output_path




def generate_architecture_png(input_json: dict, output_file="c4_ai_diagram"):
    """
        Full pipeline: JSON -> AI Layout -> PNG
        """
    print("✅ Generating AI layout plan...")
    layout_json_str = generate_ai_layout(input_json)
    print(layout_json_str)
    cleaned_layout_json = fix_ai_json(layout_json_str)
    print(cleaned_layout_json)
    print("✅ Rendering PNG...")
    render_layout_to_png(cleaned_layout_json, output_file)
    print(f"✅ C4 Container Diagram generated: {output_file}.png")


def ask_ai_for_design_doc(input_json: dict) -> str:
    """
    Ask AI to generate a professional system design document in Markdown format.
    """
    prompt = f"""
You are an highly experienced enterprise software architect and technical writer. You have vast knowledge on HLD patterns.

I will give you a JSON describing a system with:
- actors (users or external systems)
- services (microservices or internal components)
- interactions (connections between actors, services, and databases)

Generate a **professional System Design Document** in **Markdown** with the following sections:

# System Design Document

1. **System Overview**
   - One paragraph summary of the system based on the JSON

2. **High-Level Architecture**
   - List all actors, services, and databases with brief descriptions
   - Provide a bullet list of interactions
   - Provide detailed information about interactions
   - Explain the High-Level Architecture

3. **Component Responsibilities**
   - Present a table with: Component | Type | Responsibility

4. **Data Flow & Security Considerations**
   - Stepwise detailed description of data flow
   - Security, authentication, and scaling considerations

5. **Future Enhancements**
   - Suggest 2-3 improvements for scalability, maintainability, or monitoring

**Output Rules:**
- Respond in **Markdown only**
- Be **concise but professional**
- No commentary outside the document
- **Do not start with any sentences like "Here is..." or "The document is..."**

Here is the system JSON:

{json.dumps(input_json, indent=2)}
"""

    response = ai_response(prompt, "You are a helpful assistant that outputs JSON only.")

    return response.choices[0].message.content.strip()


def generate_design_doc(input_json: dict):
    """
        Generate only a professional design document from JSON.
        """
    print("✅ Generating AI Design Document...")
    markdown_doc = ask_ai_for_design_doc(input_json)

    # Save Markdown
    md_file = "Design_Document.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(markdown_doc)
    print(f"✅ Design document generated: {md_file}")
    png_file = "c4_ai_full.png"

    # Generate PDF
    docx_file = md_to_docx(md_file, png_file)

    # Step 2: Convert DOCX -> PDF
    docx_to_pdf(docx_file)


def extract_architecture_from_chunk(chunk_text):
    prompt = f"""
    You are a senior software architect.
    Analyze the following business requirement chunk and extract architecture details for High-Level Design (HLD).
    Return **ONLY a valid JSON object** following this strict schema:

    {{
      "actors": [
        {{
          "name": "string - actor name",
          "type": "External | Internal | External System"
        }}
      ],
      "microservices": [
        {{
          "name": "string - service name",
          "db": "string - database name or null if no db",
          "exposes": ["REST | gRPC | GraphQL | WebSocket | Event"],
          "consumes": ["REST | Queue | Event | DB"],
          "scaling": "string - e.g., AutoScale, Fixed, OnDemand",
          "criticality": "High | Medium | Low"
        }}
      ],
      "databases": [
        {{
          "name": "string - database name",
          "type": "SQL | NoSQL | InMemory",
          "used_by": ["list of microservice names"]
        }}
      ],
      "events": [
        {{
          "from": "string - source service or actor",
          "to": "string - destination service or actor",
          "type": "REST | Queue | Event | DB",
          "description": "string - purpose of this interaction"
        }}
      ]
    }}

    Rules:
    - JSON must be syntactically valid. Double-check for missing commas, brackets, or quotes.
    - Use double quotes for all JSON keys and values.
    - Include all services, DBs, and interactions inferred from the chunk.
    - If an element is not applicable, use null or an empty list.
    - No markdown, no explanations, no extra text outside JSON.

    Business Requirement Chunk:
    {chunk_text}
    """

    response = ai_response(prompt, "You are a helpful assistant that outputs JSON only.")

    raw_output = response.choices[0].message.content.strip()
    data_parsed = fix_ai_json(raw_output)
    print(f"parsed data_parsed: {data_parsed}")
    if not data_parsed:
        print("⚠️ JSON parse error, skipping chunk...")
    return data_parsed


def ai_response(prompt, system_role):
    response = client.chat.completions.create(
        model=MODEL_NAME,
        messages=[
            {"role": "system", "content": system_role},
            {"role": "user", "content": prompt}
        ],
        temperature=0
    )
    return response


def md_to_docx(md_file: str, png_file: str = None) -> str:
    """
    Convert a Markdown file to Word (DOCX) and embed PNG if provided.
    Returns the path to the generated DOCX file.
    """
    # Read Markdown
    with open(md_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Convert Markdown to HTML
    html_content = markdown(md_content, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html_content, "html.parser")

    # Create Word Document
    doc = Document()

    def add_list_items(parent_element, list_type="ul"):
        """Recursively add list items to the DOCX document."""
        for li in parent_element.find_all("li", recursive=False):
            # Add the current list item
            if list_type == "ul":
                doc.add_paragraph(li.text, style='List Bullet')
            else:
                doc.add_paragraph(li.text, style='List Number')

            # Check if the list item contains nested lists
            for nested_list in li.find_all(["ul", "ol"], recursive=False):
                add_list_items(nested_list, list_type=nested_list.name)

    for element in soup.children:
        if element.name in ["h1", "h2", "h3"]:
            level = int(element.name[1])
            doc.add_heading(element.text, level=level)

        elif element.name == "p":
            doc.add_paragraph(element.text)

        elif element.name == "ul":
            add_list_items(element, "ul")

        elif element.name == "ol":
            add_list_items(element, "ol")

        elif element.name == "table":
            rows = element.find_all("tr")
            if rows:
                cols = rows[0].find_all(["td", "th"])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for j, cell in enumerate(cells):
                        table.rows[i].cells[j].text = cell.text

    # Embed PNG if available
    if png_file and os.path.exists(png_file):
        doc.add_page_break()
        doc.add_heading("System Diagram", level=2)
        doc.add_picture(png_file, width=Inches(6))

    # Save DOCX
    docx_file = md_file.replace(".md", ".docx")
    doc.save(docx_file)
    print(f"✅ Word document generated: {docx_file}")
    return docx_file


def docx_to_pdf(docx_file: str):
    """
    Convert DOCX to PDF using docx2pdf.
    """
    pdf_file = docx_file.replace(".docx", ".pdf")
    convert(docx_file, pdf_file)
    print(f"✅ Final PDF generated: {pdf_file}")
    return pdf_file


def main():
    global actor, svc, inter, architecture, graphviz_path
    brd_text = load_brd(BRD_FILE)
    chunks = split_document(brd_text)
    print(f"✅ BRD loaded. Total chunks: {len(chunks)}")
    # brd_text = load_brd("business_requirements.pdf")

    # ---------------------
    # STEP 2: Extract JSON from each chunk
    # ---------------------
    all_actors = {}  # key: actor name, value: actor dict
    all_services = {}  # key: service name, value: full service dict
    all_databases = {}  # key: db name, value: db dict
    all_events = set()  # tuple (from, to, type, description)

    for idx, chunk in enumerate(chunks):
        print(f"Processing chunk {idx + 1}/{len(chunks)} ...")
        data = extract_architecture_from_chunk(chunk)  # AI JSON output

        # ---------------- Actors ----------------
        for actor in data.get("actors", []):
            if isinstance(actor, dict):
                name = actor.get("name", str(actor))
                actor_type = actor.get("type", "External")
            else:
                name = str(actor)
                actor_type = "External"
            all_actors[name] = {"name": name, "type": actor_type}

        # ---------------- Microservices ----------------
        for svc in data.get("microservices", data.get("services", [])):
            name = svc.get("name")
            if not name:
                continue
            # Merge or create service entry
            if name not in all_services:
                all_services[name] = {
                    "name": name,
                    "db": svc.get("db"),
                    "exposes": svc.get("exposes", []),
                    "consumes": svc.get("consumes", []),
                    "scaling": svc.get("scaling", "Unknown"),
                    "criticality": svc.get("criticality", "Medium")
                }

            # Register DB in all_databases if exists
            db_name = svc.get("db")
            if db_name:
                if db_name not in all_databases:
                    all_databases[db_name] = {
                        "name": db_name,
                        "type": "Unknown",
                        "used_by": [name]
                    }
                else:
                    if name not in all_databases[db_name]["used_by"]:
                        all_databases[db_name]["used_by"].append(name)

        # ---------------- Events / Interactions ----------------
        for inter in data.get("events", data.get("interactions", [])):
            src = inter.get("from")
            dst = inter.get("to")
            typ = inter.get("type", "Unknown")
            desc = inter.get("description", "")
            if src and dst:
                tup = (src, dst, typ, desc)
                all_events.add(tup)

    print("✅ Architecture extracted from all chunks!")

    # ---------------------
    # STEP 3: Merge into final JSON
    # ---------------------
    architecture = {
        "actors": list(all_actors.values()),
        "microservices": list(all_services.values()),
        "databases": list(all_databases.values()),
        "events": [
            {"from": f, "to": t, "type": typ, "description": d or ""}
            for f, t, typ, d in all_events
        ]
    }

    with open("merged_architecture.json", "w", encoding="utf-8") as f:
        json.dump(architecture, f, indent=2)

    print("✅ Merged architecture JSON saved: merged_architecture.json")

    # ---------------------
    # STEP 3: Generate C4 Container Diagram
    # ---------------------
    graphviz_path = r"C:\Graphviz-13.1.1\Graphviz-13.1.1-win64\bin"
    generate_architecture_png(architecture, output_file="c4_ai_full")

    # ---------------------
    # STEP 4: Generate Design Document
    # ---------------------
    generate_design_doc(architecture)


if __name__ == "__main__":
   
    main()




