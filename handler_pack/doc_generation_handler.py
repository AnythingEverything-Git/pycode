import os

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
from markdown import markdown

from ai_prompts import ai_for_design_doc_prompt
from ai_uitls import ai_repsonse_utility


def ask_ai_for_design_doc(input_json: dict) -> str:
    """
    Ask AI to generate a professional system design document in Markdown format.
    """
    prompt = ai_for_design_doc_prompt.ask_ai_for_design_doc(input_json)
    response = ai_repsonse_utility.ai_response(prompt, "You are a helpful assistant that outputs JSON only.")
    return response.choices[0].message.content.strip()


def generate_design_doc(input_json: dict):
    """
        Generate only a professional design document from JSON.
        """
    print("✅ Generating AI Design Document...")
    markdown_doc = ask_ai_for_design_doc(input_json)

    # Save Markdown
    md_file = "Design_Document.md"
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(markdown_doc)
    print(f"✅ Design document generated: {md_file}")
    png_file = "c4_ai_full.png"

    # Generate PDF
    docx_file = md_to_docx(md_file, png_file)

    # Step 2: Convert DOCX -> PDF
    docx_to_pdf(docx_file)


def md_to_docx(md_file: str, png_file: str = None) -> str:
    """
    Convert a Markdown file to Word (DOCX) and embed PNG if provided.
    Returns the path to the generated DOCX file.
    """
    # Read Markdown
    with open(md_file, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Convert Markdown to HTML
    html_content = markdown(md_content, extensions=["tables", "fenced_code"])
    soup = BeautifulSoup(html_content, "html.parser")

    # Create Word Document
    doc = Document()

    def add_list_items(parent_element, list_type="ul"):
        """Recursively add list items to the DOCX document."""
        for li in parent_element.find_all("li", recursive=False):
            # Add the current list item
            if list_type == "ul":
                doc.add_paragraph(li.text, style='List Bullet')
            else:
                doc.add_paragraph(li.text, style='List Number')

            # Check if the list item contains nested lists
            for nested_list in li.find_all(["ul", "ol"], recursive=False):
                add_list_items(nested_list, list_type=nested_list.name)

    for element in soup.children:
        if element.name in ["h1", "h2", "h3"]:
            level = int(element.name[1])
            doc.add_heading(element.text, level=level)

        elif element.name == "p":
            doc.add_paragraph(element.text)

        elif element.name == "ul":
            add_list_items(element, "ul")

        elif element.name == "ol":
            add_list_items(element, "ol")

        elif element.name == "table":
            rows = element.find_all("tr")
            if rows:
                cols = rows[0].find_all(["td", "th"])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                table.style = "Table Grid"
                for i, row in enumerate(rows):
                    cells = row.find_all(["td", "th"])
                    for j, cell in enumerate(cells):
                        table.rows[i].cells[j].text = cell.text

    # Embed PNG if available
    if png_file and os.path.exists(png_file):
        doc.add_page_break()
        doc.add_heading("System Diagram", level=2)
        doc.add_picture(png_file, width=Inches(6))

    # Save DOCX
    docx_file = md_file.replace(".md", ".docx")
    doc.save(docx_file)
    print(f"✅ Word document generated: {docx_file}")
    return docx_file


def docx_to_pdf(docx_file: str):
    """
    Convert DOCX to PDF using docx2pdf.
    """
    pdf_file = docx_file.replace(".docx", ".pdf")
    convert(docx_file, pdf_file)
    print(f"✅ Final PDF generated: {pdf_file}")
    return pdf_file
